import io
import re
import json
import html
from typing import List, Tuple, Dict, Set

import streamlit as st




# === Optional/standard NLP + parsing libs (all offline) ===
import nltk
from nltk.corpus import stopwords
from nltk.stem import WordNetLemmatizer

from sklearn.feature_extraction.text import CountVectorizer, TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

from PyPDF2 import PdfReader
try:
    import docx  # python-docx
except Exception:
    docx = None

# -------------------------------
# One-time NLTK setup (handles first run gracefully)
# -------------------------------
for pkg in ["punkt", "wordnet", "omw-1.4", "stopwords"]:
    try:
        nltk.data.find(f"tokenizers/{pkg}")
    except LookupError:
        nltk.download(pkg)

DEFAULT_STOPWORDS = set(stopwords.words("english"))
DEFAULT_STOPWORDS.update({
    # domain-neutral extras
    "experience", "year", "years", "etc", "using", "use", "work", "working",
    "knowledge", "ability", "skills", "requirements", "requirement", "preferred",
    "plus", "strong", "excellent", "proven", "good", "great", "must", "should",
    "including", "based", "within", "across", "across", "responsibilities",
    "role", "roles", "responsibility", "candidate", "candidates", "team",
    "teams", "stakeholders", "fit", "right", "job", "position", "apply", "applying",
})

# A small normalization + synonym map (extend as needed)
SYNONYMS = {
    "ms excel": "excel",
    "microsoft excel": "excel",
    "excel": "excel",
    "spreadsheets": "excel",
    "ml": "machine learning",
    "ai": "artificial intelligence",
    "scikit learn": "scikit-learn",
    "sklearn": "scikit-learn",
    "py torch": "pytorch",
    "tf": "tensorflow",
    "nlp": "natural language processing",
    "sql": "sql",
    "nosql": "nosql",
    "postgres": "postgresql",
    "js": "javascript",
    "ts": "typescript",
    "power bi": "powerbi",
    "git hub": "github",
    "oop": "object oriented programming",
}

# Simple curated skill-like tokens to bias extraction
CURATED_HINTS = set([
    # programming
    "python","java","c","c++","c#","javascript","typescript","go","rust","ruby","php","scala","matlab","r",
    # data/ML
    "sql","mysql","postgresql","mongodb","nosql","hive","spark","hadoop","excel","powerbi","tableau","pandas","numpy","matplotlib","scikit-learn","tensorflow","pytorch","keras","xgboost","lightgbm","nlp","computer vision","statistics","probability","data analysis","data visualization","feature engineering","model deployment",
    # devops/cloud
    "docker","kubernetes","aws","azure","gcp","ci/cd","linux","bash","git","github","gitlab",
    # soft/other
    "communication","leadership","problem solving","time management","teamwork","agile","scrum",
])

lemmatizer = WordNetLemmatizer()

# -------------------------------
# Utility functions
# -------------------------------

def normalize_phrase(s: str) -> str:
    s = s.lower().strip()
    s = re.sub(r"[\n\r]+", " ", s)
    s = re.sub(r"[^a-z0-9+#./&\- ]+", " ", s)  # keep some symbols used in skills
    s = re.sub(r"\s+", " ", s)
    # apply synonym map for multi-token phrases greedily
    for k, v in SYNONYMS.items():
        s = re.sub(rf"\b{re.escape(k)}\b", v, s)
    return s.strip()


def tokenize_words(s: str) -> List[str]:
    tokens = nltk.word_tokenize(s)
    out = []
    for t in tokens:
        t2 = t.lower()
        if t2 in DEFAULT_STOPWORDS:
            continue
        if len(t2) <= 1:
            continue
        out.append(lemmatizer.lemmatize(t2))
    return out


def read_pdf(file_obj) -> str:
    text = []
    try:
        reader = PdfReader(file_obj)
        for page in reader.pages:
            try:
                text.append(page.extract_text() or "")
            except Exception:
                pass
    except Exception:
        file_obj.seek(0)
        reader = PdfReader(io.BytesIO(file_obj.read()))
        for page in reader.pages:
            try:
                text.append(page.extract_text() or "")
            except Exception:
                pass
    return "\n".join(text)


def read_docx(file_obj) -> str:
    if docx is None:
        return ""
    try:
        file_obj.seek(0)
        document = docx.Document(file_obj)
        return "\n".join(p.text for p in document.paragraphs)
    except Exception:
        try:
            document = docx.Document(io.BytesIO(file_obj.read()))
            return "\n".join(p.text for p in document.paragraphs)
        except Exception:
            return ""


def extract_text_from_upload(uploaded_file) -> str:
    if uploaded_file is None:
        return ""
    name = (uploaded_file.name or "").lower()
    if name.endswith(".pdf"):
        return read_pdf(uploaded_file)
    elif name.endswith(".docx") or name.endswith(".doc"):
        return read_docx(uploaded_file)
    else:
        try:
            return uploaded_file.read().decode("utf-8", errors="ignore")
        except Exception:
            return ""


def top_terms_by_frequency(text: str, ngram_range=(1,3), top_k: int = 50) -> List[str]:
    norm = normalize_phrase(text)
    vectorizer = CountVectorizer(ngram_range=ngram_range, stop_words=list(DEFAULT_STOPWORDS))
    X = vectorizer.fit_transform([norm])
    freqs = X.toarray().sum(axis=0)
    terms = vectorizer.get_feature_names_out()
    pairs = sorted(zip(terms, freqs), key=lambda x: x[1], reverse=True)
    cleaned = []
    for term, _ in pairs:
        if len(term) < 2:
            continue
        if all(ch.isdigit() for ch in term):
            continue
        if term in CURATED_HINTS or (" " in term and len(term) >= 5) or term.isalpha():
            cleaned.append(term)
        if len(cleaned) >= top_k:
            break
    return cleaned


def derive_keywords_from_jd(jd_text: str, extra_hints: Set[str] = None, top_k: int = 40) -> List[str]:
    norm = normalize_phrase(jd_text)
    curated_present = [h for h in CURATED_HINTS if re.search(rf"\b{re.escape(h)}\b", norm)]
    freq_terms = top_terms_by_frequency(norm, ngram_range=(1,3), top_k=top_k)
    section_terms = []
    for line in norm.splitlines():
        if any(key in line for key in ["requirements", "requirement", "skills", "qualifications", "responsibilities"]):
            parts = re.split(r"[,;/]", line)
            for p in parts:
                p = p.strip()
                if 2 <= len(p) <= 60 and p not in DEFAULT_STOPWORDS:
                    section_terms.append(p)
    cand = curated_present + freq_terms + section_terms
    normalized: List[str] = []
    seen = set()
    for t in cand:
        t2 = normalize_phrase(t)
        toks = [w for w in t2.split() if w not in DEFAULT_STOPWORDS]
        if not toks:
            continue
        t2 = " ".join(toks)
        if len(t2) < 2:
            continue
        if t2 in seen:
            continue
        seen.add(t2)
        normalized.append(t2)
    if extra_hints:
        for h in extra_hints:
            h2 = normalize_phrase(h)
            if h2 and h2 not in seen:
                seen.add(h2)
                normalized.append(h2)
    normalized = [t for t in normalized if re.search(r"[a-z]", t)]
    return normalized


def find_matches_in_resume(keywords: List[str], resume_text: str) -> Tuple[List[str], List[str]]:
    norm_resume = normalize_phrase(resume_text)
    matched, missing = [], []
    for kw in keywords:
        kw_norm = normalize_phrase(kw)
        pattern = rf"(?<![a-z0-9]){re.escape(kw_norm)}(?![a-z0-9])"
        if re.search(pattern, norm_resume):
            matched.append(kw_norm)
        else:
            words = kw_norm.split()
            if len(words) >= 2:
                contained = sum(1 for w in words if re.search(rf"(?<![a-z0-9]){re.escape(w)}(?![a-z0-9])", norm_resume))
                if contained / len(words) >= 0.6:
                    matched.append(kw_norm)
                else:
                    missing.append(kw_norm)
            else:
                missing.append(kw_norm)

    def unique_keep(seq):
        seen = set()
        out = []
        for x in seq:
            if x not in seen:
                seen.add(x)
                out.append(x)
        return out

    return unique_keep(matched), unique_keep(missing)


def compute_cosine_similarity(jd_text: str, resume_text: str) -> float:
    tfidf = TfidfVectorizer(stop_words=list(DEFAULT_STOPWORDS))
    X = tfidf.fit_transform([jd_text, resume_text])
    sim = cosine_similarity(X[0:1], X[1:2])[0][0]
    return float(sim)


def color_badge_list(items: List[str], good: bool = True) -> str:
    color = "#16a34a" if good else "#dc2626"
    bg = "#dcfce7" if good else "#fee2e2"
    html_badges = [
        f"<span style='display:inline-block;margin:4px 6px;padding:6px 10px;"
        f"border-radius:999px;background:{bg};color:{color};border:1px solid {color}22;"
        f"font-size:0.9rem;'>{html.escape(item)}</span>"
        for item in items
    ]
    return "".join(html_badges)


def highlight_text(text: str, terms: List[str]) -> str:
    norm_text = text
    safe_terms = sorted(set(terms), key=len, reverse=True)
    for t in safe_terms:
        if not t or t.isspace():
            continue
        pattern = re.compile(rf"(?i)(?<![a-z0-9])({re.escape(t)})+(?![a-z0-9])")
        norm_text = pattern.sub(r"<mark style='background:#fff3b0;'>\\1</mark>", norm_text)
    return norm_text


# -------------------------------
# Streamlit UI
# -------------------------------

st.set_page_config(page_title="Resume Keyword Matcher (Local, No API)", page_icon="🧠", layout="wide")

st.markdown(
    """
    <style>
    .small-note { color: #475569; font-size: 0.9rem; }
    .metric-box { padding: 16px; border-radius: 14px; background: #f8fafc; border: 1px solid #e2e8f0; }
    .section { padding: 14px 16px; border-radius: 16px; background: #ffffff; border: 1px solid #e5e7eb; }
    .sec-title { font-weight: 700; font-size: 1.05rem; margin-bottom: 8px; }
    </style>
    """,
    unsafe_allow_html=True,
)

st.title("🧠 Resume Keyword Matcher — 100% Local (No API)")
st.caption("Upload your resume (PDF/DOCX/Text), paste the Job Description, and get matched vs missing skills with a match score.")

colL, colR = st.columns([1, 1])

with colL:
    uploaded_resume = st.file_uploader(
        "Upload your resume (PDF / DOCX / TXT)",
        type=["pdf", "docx", "doc", "txt"],
    )
    resume_text = extract_text_from_upload(uploaded_resume)
    if resume_text:
        st.success("Resume loaded ✅")
        with st.expander("Preview extracted resume text", expanded=False):
            st.text_area("", resume_text[:5000], height=200)

with colR:
    jd_text = st.text_area("Paste Job Description", height=260, placeholder="Paste the JD here…")

with st.expander("Optional: add custom keywords (comma separated)"):
    extra_kw_input = st.text_input("Custom keywords", value="")
    extra_hints = set([k.strip() for k in extra_kw_input.split(",") if k.strip()]) if extra_kw_input else set()

st.divider()

# Sidebar settings
with st.sidebar:
    st.header("Settings")
    ngram_max = st.slider("Max n-gram for JD keyword extraction", 1, 3, 2)
    top_k = st.slider("Max keywords to consider", 10, 150, 60)
    st.caption("Higher values may include more noise but catch more specific phrases.")

analyze = st.button("🔍 Analyze")

if analyze:
    if not resume_text:
        st.error("Please upload a resume file first.")
        st.stop()
    if not jd_text.strip():
        st.error("Please paste the Job Description.")
        st.stop()

    keywords = derive_keywords_from_jd(jd_text, extra_hints=extra_hints, top_k=top_k)

    def score_kw(k: str) -> float:
        base = len(k.split())
        if k in CURATED_HINTS:
            base += 1.5
        return base
    keywords = sorted(dict.fromkeys(keywords), key=score_kw, reverse=True)

    matched, missing = find_matches_in_resume(keywords, resume_text)

    total = max(1, len(set(keywords)))
    matched_set = set(matched)
    match_pct = round(100.0 * len(matched_set) / total, 1)

    cos_sim = compute_cosine_similarity(jd_text, resume_text)

    m1, m2, m3 = st.columns(3)
    with m1:
        st.markdown("<div class='metric-box'>", unsafe_allow_html=True)
        st.metric("Match % (keywords)", f"{match_pct}%", help="Matched keywords / total extracted keywords from JD")
        st.markdown("</div>", unsafe_allow_html=True)
    with m2:
        st.markdown("<div class='metric-box'>", unsafe_allow_html=True)
        st.metric("Cosine similarity", f"{cos_sim:.3f}", help="TF-IDF cosine similarity between JD and Resume (overall fit proxy)")
        st.markdown("</div>", unsafe_allow_html=True)
    with m3:
        st.markdown("<div class='metric-box'>", unsafe_allow_html=True)
        st.metric("Keywords considered", f"{total}")
        st.markdown("</div>", unsafe_allow_html=True)

    st.subheader("✅ Matched keywords")
    if matched:
        st.markdown(color_badge_list(matched, good=True), unsafe_allow_html=True)
    else:
        st.info("No matches found against the extracted keywords.")

    st.subheader("❌ Missing keywords")
    if missing:
        st.markdown(color_badge_list(missing, good=False), unsafe_allow_html=True)
    else:
        st.success("Looks like you covered all extracted keywords!")

    st.divider()

    c1, c2 = st.columns(2)
    with c1:
        st.markdown("<div class='section'>", unsafe_allow_html=True)
        st.markdown("<div class='sec-title'>JD with highlights</div>", unsafe_allow_html=True)
        st.markdown(highlight_text(jd_text, matched), unsafe_allow_html=True)
        st.markdown("</div>", unsafe_allow_html=True)
    with c2:
        st.markdown("<div class='section'>", unsafe_allow_html=True)
        st.markdown("<div class='sec-title'>Resume with highlights</div>", unsafe_allow_html=True)
        st.markdown(highlight_text(resume_text, matched), unsafe_allow_html=True)
        st.markdown("</div>", unsafe_allow_html=True)
    


